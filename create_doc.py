from docx import Document
from docx.shared import Cm

document = Document()
document.add_heading('CSY1020 - TCA 2 - 2023', 0)

for i in range(1,5):
    ext = "py" if i != 3 else "ipynb" 
    with open(f'Question{i}.{ext}', 'r') as file:
        document.add_paragraph(f'Question {i}', style='Heading 2')
        document.add_paragraph(file.read())
        document.add_picture(f'./screenshots/q{i}.PNG', width=Cm(8))

document.save('ayodejiOsasona.docx')
